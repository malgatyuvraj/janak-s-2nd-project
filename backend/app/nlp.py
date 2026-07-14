"""NLP utilities for the ATS platform.

Implements:
  - resume text extraction (PDF, DOCX, TXT)
  - skill extraction from free text using a curated skills dictionary
  - experience-year extraction
  - TF-IDF based matching between resume and job description
  - candidate ranking for a given job

This deliberately avoids heavy ML libraries — the math (TF-IDF + cosine
similarity) is implemented with numpy so it runs on Python 3.14+.
"""
from __future__ import annotations

import io
import re
from collections import Counter
from typing import Dict, List, Set, Tuple

import numpy as np
from PyPDF2 import PdfReader
from docx import Document


# ---------------------------------------------------------------------------
# Curated skills dictionary — covers what we see in real engineering/data
# resumes.  Matching is case-insensitive and substring-based so "React.js"
# also matches "react".
# ---------------------------------------------------------------------------
SKILL_DICTIONARY: List[str] = [
    # languages
    "python", "javascript", "typescript", "java", "kotlin", "swift",
    "c++", "c#", "go", "golang", "rust", "ruby", "php", "scala", "r",
    "sql", "html", "css",
    # frontend
    "react", "react.js", "next.js", "vue", "vue.js", "angular", "svelte",
    "redux", "tailwind", "bootstrap", "sass", "webpack", "vite",
    # backend / frameworks
    "node.js", "express", "fastapi", "django", "flask", "spring boot",
    "rails", ".net", "laravel", "nestjs",
    # data / ml
    "machine learning", "deep learning", "nlp", "natural language processing",
    "computer vision", "tensorflow", "pytorch", "keras", "scikit-learn",
    "pandas", "numpy", "matplotlib", "data analysis", "data science",
    # databases
    "postgresql", "mysql", "mongodb", "redis", "elasticsearch", "sqlite",
    "dynamodb", "cassandra", "oracle",
    # cloud / devops
    "aws", "azure", "gcp", "google cloud", "docker", "kubernetes",
    "terraform", "ansible", "jenkins", "github actions", "ci/cd",
    # mobile
    "ios", "android", "react native", "flutter", "xamarin",
    # tools / methodologies
    "git", "rest api", "graphql", "microservices", "agile", "scrum",
    "jira", "figma", "linux", "bash",
]


# Build a lookup: normalized skill -> display name
_SKILL_LOOKUP: Dict[str, str] = {
    s.lower().strip(): s for s in SKILL_DICTIONARY
}


# ---------------------------------------------------------------------------
# Text extraction
# ---------------------------------------------------------------------------
def extract_text_from_pdf(data: bytes) -> str:
    reader = PdfReader(io.BytesIO(data))
    parts = []
    for page in reader.pages:
        try:
            parts.append(page.extract_text() or "")
        except Exception:
            parts.append("")
    return "\n".join(parts)


def extract_text_from_docx(data: bytes) -> str:
    doc = Document(io.BytesIO(data))
    parts = [p.text for p in doc.paragraphs]
    # also walk tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                parts.append(cell.text)
    return "\n".join(parts)


def extract_text_from_upload(filename: str, content: bytes) -> str:
    """Dispatch on file extension. Plain text fallback for unknown types."""
    name = (filename or "").lower()
    if name.endswith(".pdf"):
        return extract_text_from_pdf(content)
    if name.endswith(".docx"):
        return extract_text_from_docx(content)
    # treat as plain text for .txt or unknown
    try:
        return content.decode("utf-8", errors="ignore")
    except Exception:
        return ""


# ---------------------------------------------------------------------------
# Cleaning & tokenization
# ---------------------------------------------------------------------------
_TOKEN_RE = re.compile(r"[a-zA-Z][a-zA-Z0-9.+#-]*")


def clean_text(text: str) -> str:
    text = (text or "").lower()
    text = re.sub(r"https?://\S+", " ", text)
    text = re.sub(r"[^a-z0-9\s.+#-]", " ", text)
    text = re.sub(r"\s+", " ", text)
    return text.strip()


def tokenize(text: str) -> List[str]:
    return _TOKEN_RE.findall((text or "").lower())


# A small English stop-word list — keeps the implementation self-contained.
STOPWORDS: Set[str] = {
    "a", "an", "the", "and", "or", "but", "if", "then", "is", "are", "was",
    "were", "be", "been", "being", "have", "has", "had", "do", "does", "did",
    "i", "you", "he", "she", "it", "we", "they", "me", "him", "her", "us",
    "them", "my", "your", "his", "its", "our", "their", "this", "that",
    "these", "those", "of", "in", "on", "at", "to", "from", "by", "with",
    "as", "for", "about", "into", "over", "under", "up", "down", "out",
    "off", "no", "not", "yes", "can", "could", "should", "would", "will",
    "shall", "may", "might", "must", "via", "per", "etc", "than", "so",
    "also", "just", "very", "more", "less", "most", "least", "some",
    "any", "all", "each", "every", "other", "such", "only", "own", "same",
    "too", "now", "here", "there", "when", "where", "why", "how", "what",
    "who", "whom", "which", "while", "during", "before", "after",
    "between", "through", "above", "below", "up", "again", "further",
    "once", "am", "pm",
}


def meaningful_tokens(text: str) -> List[str]:
    return [t for t in tokenize(text) if t not in STOPWORDS and len(t) > 1]


# ---------------------------------------------------------------------------
# Skill extraction
# ---------------------------------------------------------------------------
def extract_skills(text: str) -> List[str]:
    """Return the list of canonical skills present in `text`.

    Uses case-insensitive substring matching against SKILL_DICTIONARY.  Multi-
    word skills (e.g. "machine learning") are matched first so they win over
    overlapping single-word tokens.
    """
    if not text:
        return []
    lower = text.lower()
    found: List[str] = []
    seen: Set[str] = set()

    # match multi-word skills first (sorted by length desc) so longer phrases
    # are preferred over their components
    multi_word = sorted(
        [s for s in _SKILL_LOOKUP if " " in s or "." in s or "#" in s],
        key=len,
        reverse=True,
    )
    for skill in multi_word:
        # use word-ish boundaries; for skills with dots/hashes, just substring
        if "." in skill or "#" in skill or skill.endswith("+"):
            pattern = re.escape(skill)
        else:
            pattern = r"\b" + re.escape(skill) + r"\b"
        if re.search(pattern, lower) and skill not in seen:
            found.append(_SKILL_LOOKUP[skill])
            seen.add(skill)

    # then single-word skills
    for skill in _SKILL_LOOKUP:
        if " " in skill or "." in skill or "#" in skill or skill.endswith("+"):
            continue
        pattern = r"\b" + re.escape(skill) + r"\b"
        if re.search(pattern, lower) and skill not in seen:
            found.append(_SKILL_LOOKUP[skill])
            seen.add(skill)

    return found


# ---------------------------------------------------------------------------
# Experience extraction
# ---------------------------------------------------------------------------
_EXPERIENCE_RE = re.compile(
    r"(\d{1,2})\s*\+?\s*(?:years?|yrs?)\s*(?:of)?\s*(?:experience|exp)?",
    re.IGNORECASE,
)


def extract_experience_years(text: str) -> int:
    """Return the largest experience figure mentioned in the resume."""
    if not text:
        return 0
    matches = _EXPERIENCE_RE.findall(text)
    nums = [int(m) for m in matches if m.isdigit() and 0 < int(m) < 50]
    return max(nums) if nums else 0


# ---------------------------------------------------------------------------
# TF-IDF + cosine similarity (numpy only)
# ---------------------------------------------------------------------------
def _build_vocab(docs: List[List[str]]) -> Dict[str, int]:
    vocab: Dict[str, int] = {}
    for doc in docs:
        for tok in doc:
            if tok not in vocab:
                vocab[tok] = len(vocab)
    return vocab


def _tfidf_matrix(docs: List[List[str]]) -> np.ndarray:
    """Compute a TF-IDF matrix (rows = docs) using log-smoothed IDF."""
    vocab = _build_vocab(docs)
    n_docs = len(docs)
    n_terms = len(vocab)
    if n_terms == 0:
        return np.zeros((n_docs, 0))

    # document frequencies
    df = np.zeros(n_terms, dtype=np.float64)
    for doc in docs:
        for tok in set(doc):
            if tok in vocab:
                df[vocab[tok]] += 1.0

    # smoothed IDF: log((1 + N) / (1 + df)) + 1
    idf = np.log((1.0 + n_docs) / (1.0 + df)) + 1.0

    mat = np.zeros((n_docs, n_terms), dtype=np.float64)
    for i, doc in enumerate(docs):
        if not doc:
            continue
        counts = Counter(doc)
        total = sum(counts.values())
        for tok, c in counts.items():
            j = vocab.get(tok)
            if j is None:
                continue
            tf = c / total
            mat[i, j] = tf * idf[j]
    # L2-normalize each row for cosine similarity via dot product
    norms = np.linalg.norm(mat, axis=1, keepdims=True)
    norms[norms == 0] = 1.0
    return mat / norms


def cosine_similarity(a: np.ndarray, b: np.ndarray) -> np.ndarray:
    """Cosine similarity between row vectors. Inputs are assumed L2-normalized."""
    if a.size == 0 or b.size == 0:
        return np.zeros((a.shape[0], b.shape[0]))
    return a @ b.T


# ---------------------------------------------------------------------------
# High-level matching
# ---------------------------------------------------------------------------
def compute_match(job_description: str,
                  job_required_skills: List[str],
                  resume_text: str,
                  resume_skills: List[str]) -> Dict:
    """Compute a 0..100 match score between a resume and a job.

    The score blends:
      - TF-IDF cosine similarity between full description and full resume (40%)
      - Overlap between the resume's skills and the job's required skills (60%)
    """
    jd_tokens = meaningful_tokens(job_description or "")
    resume_tokens = meaningful_tokens(resume_text or "")

    jd_skills_norm = {s.lower().strip() for s in (job_required_skills or [])}
    resume_skills_norm = {s.lower().strip() for s in (resume_skills or [])}

    if jd_tokens and resume_tokens:
        vecs = _tfidf_matrix([resume_tokens, jd_tokens])
        sim = float(cosine_similarity(vecs[0:1], vecs[1:2])[0, 0])
        sim_score = max(0.0, min(1.0, sim)) * 100.0
    else:
        sim_score = 0.0

    if jd_skills_norm:
        matched = jd_skills_norm & resume_skills_norm
        missing = jd_skills_norm - resume_skills_norm
        skill_score = (len(matched) / len(jd_skills_norm)) * 100.0
        matched_display = [
            s for s in (job_required_skills or []) if s.lower().strip() in matched
        ]
        missing_display = [
            s for s in (job_required_skills or []) if s.lower().strip() in missing
        ]
    else:
        # no explicit skill list — use skills found in description text
        jd_skills_from_text = {s.lower() for s in extract_skills(job_description or "")}
        if jd_skills_from_text:
            matched = jd_skills_from_text & resume_skills_norm
            missing = jd_skills_from_text - resume_skills_norm
            skill_score = (len(matched) / len(jd_skills_from_text)) * 100.0
            matched_display = sorted(matched)
            missing_display = sorted(missing)
        else:
            skill_score = 0.0
            matched_display = []
            missing_display = []

    final = 0.4 * sim_score + 0.6 * skill_score
    final = round(max(0.0, min(100.0, final)), 2)

    return {
        "score": final,
        "tfidf_similarity": round(sim_score, 2),
        "skill_overlap_pct": round(skill_score, 2),
        "matched_skills": matched_display,
        "missing_skills": missing_display,
    }


def rank_candidates(job_description: str,
                    job_required_skills: List[str],
                    candidates: List[Dict]) -> List[Dict]:
    """Return candidates sorted by descending match score, with scores
    attached. Each candidate dict must have 'resume_text' and 'skills'."""
    results = []
    for cand in candidates:
        cand_skills = cand.get("skills") or []
        if isinstance(cand_skills, str):
            cand_skills = [s.strip() for s in cand_skills.split(",") if s.strip()]
        m = compute_match(
            job_description,
            job_required_skills,
            cand.get("resume_text", "") or "",
            cand_skills,
        )
        out = dict(cand)
        out["match_score"] = m["score"]
        out["matched_skills"] = m["matched_skills"]
        out["missing_skills"] = m["missing_skills"]
        results.append(out)
    results.sort(key=lambda x: x["match_score"], reverse=True)
    return results


def parse_resume(filename: str, content: bytes) -> Dict:
    """One-shot resume parser: text → skills, experience, raw text."""
    text = extract_text_from_upload(filename, content)
    return {
        "resume_text": text,
        "skills": extract_skills(text),
        "experience_years": extract_experience_years(text),
    }
